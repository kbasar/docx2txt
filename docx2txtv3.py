'''This script is to copy text from documents (docx) to simple text file

'''
import shutil
import sys
import ntpath
import os
from docx import Document
# Ask user to enter docx source path
user_input_docx_path = input("Enter the path where docx files are located: ")
assert os.path.exists(user_input_docx_path), "I did not find the file at, " + str(user_input_docx_path)

user_input_txt_path = input("Enter the path where text files should be saved: ")
if os.path.exists(user_input_txt_path):
    shutil.rmtree(user_input_txt_path)
os.makedirs(user_input_txt_path)

docpath = user_input_docx_path
txtpath = user_input_txt_path

#docpath = os.path.abspath(r'C:\Users\Khairul Basar\Documents\CWD Projects\Searchline Database\00_WORKING\WL_SLOT1_submission_date_30-03-2018\1-100')
#txtpath = os.path.abspath(r'C:\Users\Khairul Basar\Documents\CWD Projects\Searchline Database\00_WORKING\WL_SLOT1_submission_date_30-03-2018\Textfiles')

for filename in os.listdir(docpath):
    document = Document(os.path.join(docpath, filename))
    # print(document.paragraphs)
    print(filename)
    savetxt = os.path.join(txtpath, ntpath.basename(filename).split('.')[0] + ".txt")
    print('Reading ' + filename)
    # print(savetxt)
    fullText = []
    for para in document.paragraphs:
        # print(para.text)
        fullText.append(para.text)
    with open(savetxt, 'w', encoding='utf-8') as newfile:
        for item in fullText:
            newfile.write("%s\n" % item)


print("Done converting Docx to txt")
