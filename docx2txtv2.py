'''This script is to copy text from documents (docx) to simple text file

'''

import sys
import ntpath
import os
from docx import Document

docpath = os.path.abspath(r'C:\Users\Khairul Basar\Documents\CWD Projects\abc1-100')
txtpath = os.path.abspath(r'C:\Users\Khairul Basar\Documents\CWD Projects\Textfiles')

for filename in os.listdir(docpath):
    document = Document(os.path.join(docpath, filename))
    # print(document.paragraphs)
    print(filename)
    savetxt = os.path.join(txtpath, ntpath.basename(filename).split('.')[0] + ".txt")
    print('Reading ' + filename)
    # print(savetxt)
    fullText = []
    for para in document.paragraphs:
        # print(para.text)
        fullText.append(para.text)
    with open(savetxt, 'w', encoding='utf-8') as newfile:
        for item in fullText:
            newfile.write("%s\n" % item)


print("Done converting Docx to txt")
